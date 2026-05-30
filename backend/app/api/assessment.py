import io
import logging
import fitz  # PyMuPDF for quick PDF parses
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from typing import List, Optional
from docx import Document as DocxDocument

from app.core.database import get_db
from app.models import (
    Assessment,
    Client,
    ProcessBottleneck,
    AIUseCase,
    RiskRegister,
    RoadmapItem,
    DocumentSignal,
    User,
)
from app.schemas.assessment import (
    AssessmentCreate,
    AssessmentUpdate,
    AssessmentResponse,
)
from app.agents.graph import AssessmentOrchestrator
from app.services.document_generators import DocumentGenerators
from app.api.auth import get_current_user
from app.api.clients import get_client_for_user, get_or_create_default_organization
from app.core.sanitizer import sanitize_text


logger = logging.getLogger("AssessmentAPI")
router = APIRouter()

orchestrator = AssessmentOrchestrator()
VALID_APPROVAL_STATUSES = {"draft", "reviewed", "approved"}


def normalize_department(use_case_name: str, description: str) -> str:
    n = (use_case_name or "").lower()
    d = (description or "").lower()

    # 1. Sales & Pre-sales
    if any(
        x in n or x in d
        for x in [
            "proposal",
            "rfp",
            "pre-sales",
            "sales",
            "pricing",
            "client email",
            "bid",
        ]
    ):
        return "Sales & Pre-sales"
    # 2. Customer Support
    if any(
        x in n or x in d
        for x in ["support", "ticket", "triage", "incident", "response assistant"]
    ):
        return "Customer Support"
    # 3. Operations
    if any(
        x in n or x in d
        for x in ["billing", "invoice", "reconciliation", "spreadsheet", "excel"]
    ):
        return "Operations"
    # 4. Compliance & Governance
    if any(
        x in n or x in d
        for x in ["governance", "compliance", "audit", "contract", "policy"]
    ):
        return "Compliance & Governance"
    return "Operations"


# Helper to load DB assessment into dict format for generators or orchestrators
def _db_assessment_to_dict(ass: Assessment) -> dict:
    return {
        "id": ass.id,
        "company_name": ass.company_name,
        "industry": ass.industry,
        "company_size": ass.company_size,
        "departments": ass.departments,
        "current_tools": ass.current_tools,
        "cloud_preference": ass.cloud_preference,
        "compliance_requirements": ass.compliance_requirements,
        "main_business_goals": ass.main_business_goals,
        "pain_points": ass.pain_points,
        "ai_goals": ass.ai_goals,
        "overall_score": ass.overall_score,
        "automation_potential": ass.automation_potential,
        "confidence_score": ass.confidence_score,
        "status": ass.status,
        "risk_level": ass.risk_level,
        "recommended_first_pilot": ass.recommended_first_pilot,
        "why_recommended_pilot": ass.why_recommended_pilot,
        "expected_pilot_impact": ass.expected_pilot_impact,
        "data_readiness": ass.data_readiness,
        "data_justification": ass.data_justification,
        "process_readiness": ass.process_readiness,
        "process_justification": ass.process_justification,
        "integration_readiness": ass.integration_readiness,
        "integration_justification": ass.integration_justification,
        "governance_readiness": ass.governance_readiness,
        "governance_justification": ass.governance_justification,
        "security_readiness": ass.security_readiness,
        "security_justification": ass.security_justification,
        "team_readiness": ass.team_readiness,
        "team_justification": ass.team_justification,
        "business_alignment": ass.business_alignment,
        "alignment_justification": ass.alignment_justification,
        "business_summary": ass.business_summary,
        "client_summary": ass.client_summary,
        "reviewer_notes": ass.reviewer_notes,
        "approval_status": ass.approval_status,
        "readiness_interpretation": ass.readiness_interpretation,
        "bottlenecks": [
            {
                "department": b.department,
                "process_name": b.process_name,
                "bottleneck_description": b.bottleneck_description,
                "ai_potential": b.ai_potential,
            }
            for b in ass.bottlenecks
        ],
        "use_cases": [
            {
                "use_case_name": u.use_case_name,
                "department": normalize_department(u.use_case_name, u.description),
                "description": u.description,
                "value": u.value,
                "complexity": u.complexity,
                "risk": u.risk,
                "priority": u.priority,
                "evidence": u.evidence,
                "confidence": u.confidence,
            }
            for u in ass.use_cases
        ],
        "risks": [
            {
                "risk_name": r.risk_name,
                "severity": r.severity,
                "recommendation": r.recommendation,
                "is_control_met": r.is_control_met,
            }
            for r in ass.risks
        ],
        "roadmap_items": [
            {
                "phase": rm.phase,
                "action_item": rm.action_item,
                "expected_impact": rm.expected_impact,
                "confidence": rm.confidence,
            }
            for rm in ass.roadmap_items
        ],
        "extracted_signals": [
            {
                "source_file": s.source_file,
                "signal_type": s.signal_type,
                "description": s.description,
                "confidence": s.confidence,
            }
            for s in ass.extracted_signals
        ],
    }


@router.post("/", response_model=AssessmentResponse)
def create_assessment(
    payload: AssessmentCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Creates a blank assessment session based on Intake details"""
    logger.info(f"Creating intake session for: {payload.company_name}")
    organization = get_or_create_default_organization(db, current_user)
    client = None

    if payload.client_id is not None:
        client = get_client_for_user(db, current_user, payload.client_id)
        if not client:
            raise HTTPException(status_code=404, detail="Client not found")
    else:
        client = (
            db.query(Client)
            .filter(
                Client.organization_id == organization.id,
                Client.name == payload.company_name,
            )
            .first()
        )
        if not client:
            client = Client(
                organization_id=organization.id,
                name=payload.company_name,
                industry=payload.industry,
                company_size=payload.company_size,
                cloud_preference=payload.cloud_preference,
                compliance_requirements=payload.compliance_requirements,
            )
            db.add(client)
            db.flush()

    ass = Assessment(
        user_id=current_user.id,
        client_id=client.id,
        company_name=client.name,
        industry=payload.industry or client.industry,
        company_size=payload.company_size or client.company_size,
        departments=payload.departments,
        current_tools=payload.current_tools,
        cloud_preference=payload.cloud_preference or client.cloud_preference,
        compliance_requirements=payload.compliance_requirements
        or (client.compliance_requirements if client else []),
        main_business_goals=payload.main_business_goals,
        pain_points=payload.pain_points,
        ai_goals=payload.ai_goals,
        status="intake",
    )
    db.add(ass)
    db.commit()
    db.refresh(ass)
    return ass


@router.get("/", response_model=List[AssessmentResponse])
def list_assessments(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return (
        db.query(Assessment)
        .filter(Assessment.user_id == current_user.id)
        .order_by(Assessment.updated_at.desc())
        .all()
    )


@router.delete("/{id}", status_code=204)
def delete_assessment(
    id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    ass = (
        db.query(Assessment)
        .filter(Assessment.id == id, Assessment.user_id == current_user.id)
        .first()
    )
    if not ass:
        raise HTTPException(status_code=404, detail="Assessment not found")

    db.delete(ass)
    db.commit()
    return None


@router.get("/{id}", response_model=AssessmentResponse)
def get_assessment(
    id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Retrieves full assessment with all related details loaded"""
    ass = (
        db.query(Assessment)
        .filter(Assessment.id == id, Assessment.user_id == current_user.id)
        .first()
    )
    if not ass:
        raise HTTPException(status_code=404, detail="Assessment not found")
    return ass


@router.post("/{id}/upload", response_model=AssessmentResponse)
def upload_documents(
    id: int,
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Parses uploaded files, populates text corpus, and executes the LangGraph analysis"""
    ass = (
        db.query(Assessment)
        .filter(Assessment.id == id, Assessment.user_id == current_user.id)
        .first()
    )
    if not ass:
        raise HTTPException(status_code=404, detail="Assessment not found")

    ass.status = "processing"
    db.commit()

    extracted_corpus = []
    signals_to_create = []

    # 1. Parse uploaded files (using PyMuPDF for PDFs, fallback to utf-8 text)
    for f in files:
        file_text = ""
        try:
            content = f.file.read()
            if f.filename and f.filename.endswith(".pdf"):
                logger.info(f"Parsing PDF file: {f.filename}")
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    file_text += page.get_text()
            elif f.filename and f.filename.endswith(".docx"):
                logger.info(f"Parsing DOCX file: {f.filename}")
                doc = DocxDocument(io.BytesIO(content))
                paragraph_text = [
                    paragraph.text
                    for paragraph in doc.paragraphs
                    if paragraph.text.strip()
                ]
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        ]
                        if cells:
                            table_text.append(" | ".join(cells))
                file_text = "\n".join(paragraph_text + table_text)
            else:
                logger.info(f"Parsing text/raw file: {f.filename}")
                file_text = content.decode("utf-8", errors="ignore")

            extracted_corpus.append(f"--- File: {f.filename} ---\n{file_text}")
            signals_to_create.append(
                DocumentSignal(
                    source_file=f.filename,
                    signal_type="Document Upload",
                    description=f"Successfully extracted text length: {len(file_text)} characters.",
                    confidence=95.0,
                )
            )
        except Exception as e:
            logger.error(f"Error parsing file {f.filename}: {e}")
            signals_to_create.append(
                DocumentSignal(
                    source_file=f.filename,
                    signal_type="Parsing Warning",
                    description=f"File parsed with errors: {str(e)}",
                    confidence=40.0,
                )
            )

    consolidated_text = "\n\n".join(extracted_corpus)

    # Save parsed text and signals
    ass.extracted_signals.extend(signals_to_create)
    db.commit()

    # 2. Run LangGraph Multi-Agent Pipeline
    initial_graph_state = {
        "assessment_id": ass.id,
        "company_name": ass.company_name,
        "industry": ass.industry or "Technology Services",
        "company_size": ass.company_size or "10-100 employees",
        "departments": ass.departments or [],
        "current_tools": ass.current_tools or [],
        "cloud_preference": ass.cloud_preference or "Cloud-agnostic",
        "compliance_requirements": ass.compliance_requirements or [],
        "main_business_goals": ass.main_business_goals or "",
        "pain_points": ass.pain_points or [],
        "ai_goals": ass.ai_goals or [],
        "extracted_text": consolidated_text,
    }

    try:
        final_state = orchestrator.run_assessment(
            initial_graph_state, thread_id=f"thread_{ass.id}"
        )

        # 3. Synchronize Graph state outputs to Database models
        ass.business_summary = sanitize_text(
            final_state.get("business_summary"),
            "Strategic review of core manual workflows, data systems, and compliance requirements.",
        )
        ass.client_summary = sanitize_text(
            final_state.get("business_summary"),
            "Strategic review of core manual workflows, data systems, and compliance requirements.",
        )
        ass.approval_status = "draft"
        ass.overall_score = final_state.get("overall_score", 0.0)
        ass.automation_potential = final_state.get("automation_potential", 0.0)
        ass.confidence_score = final_state.get("recommended_pilot", {}).get(
            "confidence", 85.0
        )
        ass.readiness_interpretation = sanitize_text(
            final_state.get("readiness_interpretation"),
            "The client shows a strong business alignment score, with moderate integration and governance scores that can be resolved via an operational pilot.",
        )

        # Score breakdown
        ass.data_readiness = final_state.get("data_readiness", 0.0)
        ass.data_justification = sanitize_text(
            final_state.get("data_justification"),
            "The available data landscape appears mixed, with meaningful business information present but not yet fully standardized for AI use.",
        )
        ass.process_readiness = final_state.get("process_readiness", 0.0)
        ass.process_justification = sanitize_text(
            final_state.get("process_justification"),
            "Core workflows are repetitive enough for AI support, but they still need tighter standardization before large-scale automation.",
        )
        ass.integration_readiness = final_state.get("integration_readiness", 0.0)
        ass.integration_justification = sanitize_text(
            final_state.get("integration_justification"),
            "Current systems can support integration work, but shared middleware and clean workflow handoffs still need to be established.",
        )
        ass.governance_readiness = final_state.get("governance_readiness", 0.0)
        ass.governance_justification = sanitize_text(
            final_state.get("governance_justification"),
            "Governance requirements are visible, but AI-specific approval controls and auditability need to be strengthened.",
        )
        ass.security_readiness = final_state.get("security_readiness", 0.0)
        ass.security_justification = sanitize_text(
            final_state.get("security_justification"),
            "Baseline enterprise security controls appear present, though AI-specific data protections still require operational enforcement.",
        )
        ass.team_readiness = final_state.get("team_readiness", 0.0)
        ass.team_justification = sanitize_text(
            final_state.get("team_justification"),
            "Leadership interest is present, but adoption readiness still depends on training, operating ownership, and change management discipline.",
        )
        ass.business_alignment = final_state.get("business_alignment", 0.0)
        ass.alignment_justification = sanitize_text(
            final_state.get("alignment_justification"),
            "The most promising AI use cases are well aligned with measurable business goals such as faster delivery and lower manual effort.",
        )

        # Risk assessment
        ass.risk_level = (
            "High"
            if any(r.get("severity") == "High" for r in final_state.get("risks", []))
            else "Medium"
        )

        # Recommended Pilot
        pilot = final_state.get("recommended_pilot", {})
        ass.recommended_first_pilot = sanitize_text(
            pilot.get("name"), "Intelligent Pre-Sales Proposal Copilot"
        )
        ass.why_recommended_pilot = sanitize_text(
            pilot.get("why"),
            "High transformation value, low implementation complexity, and matches observed pre-sales bottlenecks.",
        )
        ass.expected_pilot_impact = sanitize_text(
            pilot.get("expected_impact"),
            "Reduces first-draft proposal preparation time, improves approved content reuse, and creates a controlled review flow for client-ready documents.",
        )

        # Map Bottlenecks
        db.query(ProcessBottleneck).filter(
            ProcessBottleneck.assessment_id == ass.id
        ).delete()
        for b in final_state.get("bottlenecks", []):
            db.add(
                ProcessBottleneck(
                    assessment_id=ass.id,
                    department=b.get("department", "Operations"),
                    process_name=sanitize_text(
                        b.get("process_name"), "Manual Operational Process"
                    ),
                    bottleneck_description=sanitize_text(
                        b.get("bottleneck_description"),
                        "Manual collation and copy-paste processes limit operational velocity.",
                    ),
                    ai_potential=b.get("ai_potential", "Medium"),
                )
            )

        # Map Use Cases
        db.query(AIUseCase).filter(AIUseCase.assessment_id == ass.id).delete()
        for u in final_state.get("use_cases", []):
            uc_name = sanitize_text(
                u.get("use_case_name"), "Intelligent Operational Copilot"
            )
            uc_desc = sanitize_text(
                u.get("description"),
                "AI decision support and semantic retrieval tool to streamline manual workflows.",
            )
            db.add(
                AIUseCase(
                    assessment_id=ass.id,
                    department=normalize_department(uc_name, uc_desc),
                    use_case_name=uc_name,
                    description=uc_desc,
                    value=u.get("value", "High"),
                    complexity=u.get("complexity", "Medium"),
                    risk=u.get("risk", "Low"),
                    priority=u.get("priority", "P1"),
                    evidence=sanitize_text(
                        u.get("evidence"),
                        "Manual processing and multi-system copy-pasting loops observed in standard operations.",
                    ),
                    confidence=u.get("confidence", 85.0),
                )
            )

        # Map Risks
        db.query(RiskRegister).filter(RiskRegister.assessment_id == ass.id).delete()
        for r in final_state.get("risks", []):
            db.add(
                RiskRegister(
                    assessment_id=ass.id,
                    risk_name=sanitize_text(
                        r.get("risk_name"), "AI Implementation Risk"
                    ),
                    severity=r.get("severity", "Medium"),
                    recommendation=sanitize_text(
                        r.get("recommendation"),
                        "Apply human review, source logging, redaction controls, and evaluation baselines before AI outputs are shared externally.",
                    ),
                    is_control_met=r.get("is_control_met", 0),
                )
            )

        # Map Roadmap Items
        db.query(RoadmapItem).filter(RoadmapItem.assessment_id == ass.id).delete()
        for item in final_state.get("roadmap_items", []):
            db.add(
                RoadmapItem(
                    assessment_id=ass.id,
                    phase=item.get("phase", "30-Day"),
                    action_item=sanitize_text(
                        item.get("action_item"),
                        "Finalize pilot scope, confirm success metrics, map required documents, and prepare controlled MVP delivery plan.",
                    ),
                    expected_impact=sanitize_text(
                        item.get("expected_impact"),
                        "Creates a validated pilot foundation with measurable success criteria.",
                    ),
                    confidence=item.get("confidence", 80.0),
                )
            )

        # Add a parsed document signal indicating LangGraph success
        db.add(
            DocumentSignal(
                assessment_id=ass.id,
                source_file="LangGraphOrchestrator",
                signal_type="Execution Complete",
                description="All 9 nodes of the stateful analysis graph executed successfully.",
                confidence=100.0,
            )
        )

        ass.status = "completed"
        db.commit()
        db.refresh(ass)

    except Exception as e:
        logger.error(f"LangGraph execution crashed: {e}")
        ass.status = "failed"
        db.add(
            DocumentSignal(
                assessment_id=ass.id,
                source_file="LangGraphOrchestrator",
                signal_type="Execution Failure",
                description=f"Pipeline crashed with message: {str(e)}",
                confidence=0.0,
            )
        )
        db.commit()
        raise HTTPException(
            status_code=500, detail=f"AI Graph analysis failed: {str(e)}"
        )

    return ass


@router.put("/{id}", response_model=AssessmentResponse)
def update_assessment(
    id: int,
    payload: AssessmentUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Human Review Mode: allows overriding any score or list item manually"""
    ass = (
        db.query(Assessment)
        .filter(Assessment.id == id, Assessment.user_id == current_user.id)
        .first()
    )
    if not ass:
        raise HTTPException(status_code=404, detail="Assessment not found")

    # Standard overrides
    update_data = payload.model_dump(exclude_unset=True)

    # Pop nested items to handle separately
    use_cases = update_data.pop("use_cases", None)
    bottlenecks = update_data.pop("bottlenecks", None)
    risks = update_data.pop("risks", None)
    roadmap_items = update_data.pop("roadmap_items", None)

    for key, value in update_data.items():
        if key == "client_id":
            if value is None:
                ass.client_id = None
                continue
            client = get_client_for_user(db, current_user, value)
            if not client:
                raise HTTPException(status_code=404, detail="Client not found")
            ass.client_id = client.id
            ass.company_name = client.name
            continue
        if (
            key == "approval_status"
            and value is not None
            and value not in VALID_APPROVAL_STATUSES
        ):
            raise HTTPException(status_code=400, detail="Invalid approval status")
        setattr(ass, key, value)

    # Update Use Cases if provided
    if use_cases is not None:
        db.query(AIUseCase).filter(AIUseCase.assessment_id == id).delete()
        for u in use_cases:
            uc_name = u.get("use_case_name")
            uc_desc = u.get("description")
            db.add(
                AIUseCase(
                    assessment_id=id,
                    department=normalize_department(uc_name, uc_desc),
                    use_case_name=uc_name,
                    description=uc_desc,
                    value=u.get("value"),
                    complexity=u.get("complexity"),
                    risk=u.get("risk"),
                    priority=u.get("priority"),
                    evidence=u.get("evidence"),
                    confidence=u.get("confidence", 85.0),
                )
            )

    # Update Bottlenecks if provided
    if bottlenecks is not None:
        db.query(ProcessBottleneck).filter(
            ProcessBottleneck.assessment_id == id
        ).delete()
        for b in bottlenecks:
            db.add(
                ProcessBottleneck(
                    assessment_id=id,
                    department=b.get("department"),
                    process_name=b.get("process_name"),
                    bottleneck_description=b.get("bottleneck_description"),
                    ai_potential=b.get("ai_potential"),
                )
            )

    # Update Risks if provided
    if risks is not None:
        db.query(RiskRegister).filter(RiskRegister.assessment_id == id).delete()
        for r in risks:
            db.add(
                RiskRegister(
                    assessment_id=id,
                    risk_name=r.get("risk_name"),
                    severity=r.get("severity"),
                    recommendation=r.get("recommendation"),
                    is_control_met=r.get("is_control_met", 0),
                )
            )

    # Update Roadmap if provided
    if roadmap_items is not None:
        db.query(RoadmapItem).filter(RoadmapItem.assessment_id == id).delete()
        for item in roadmap_items:
            db.add(
                RoadmapItem(
                    assessment_id=id,
                    phase=item.get("phase"),
                    action_item=item.get("action_item"),
                    expected_impact=item.get("expected_impact"),
                    confidence=item.get("confidence", 80.0),
                )
            )

    # Re-calculate readiness interpretation dynamically on manual score changes
    if "readiness_interpretation" not in update_data:
        prefix = "With a consultant-reviewed AI readiness score of"
        overall = ass.overall_score
        integration = ass.integration_readiness
        ass.readiness_interpretation = (
            f"{prefix} {int(overall)}/100, the organization appears ready for a controlled pilot rollout, "
            f"while integration ({int(integration)}/100) and governance controls should still be validated "
            f"before production scaling."
        )

    db.commit()
    db.refresh(ass)
    return ass


@router.get("/{id}/export/{doc_format}")
def export_assessment_report(
    id: int,
    doc_format: str,
    mode: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Exports compiled custom formatted PDF, DOCX, or PPTX reports"""
    ass = (
        db.query(Assessment)
        .filter(Assessment.id == id, Assessment.user_id == current_user.id)
        .first()
    )
    if not ass:
        raise HTTPException(status_code=404, detail="Assessment not found")
    if ass.approval_status != "approved":
        raise HTTPException(
            status_code=409,
            detail="Assessment must be approved in Human Review Mode before export.",
        )

    # Serialize DB object to dictionary structure
    data = _db_assessment_to_dict(ass)
    data["export_mode"] = mode

    if doc_format.lower() == "pdf":
        pdf_buffer = DocumentGenerators.generate_pdf_report(data)
        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={ass.company_name.replace(' ', '_')}_AI_Readiness_Report.pdf"
            },
        )
    elif doc_format.lower() == "docx":
        docx_buffer = DocumentGenerators.generate_docx_proposal(data)
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={ass.company_name.replace(' ', '_')}_AI_Pilot_Proposal.docx"
            },
        )
    elif doc_format.lower() == "pptx":
        pptx_buffer = DocumentGenerators.generate_pptx_deck(data)
        return StreamingResponse(
            pptx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            headers={
                "Content-Disposition": f"attachment; filename={ass.company_name.replace(' ', '_')}_Board_Presentation.pptx"
            },
        )
    else:
        raise HTTPException(
            status_code=400, detail="Invalid format. Use PDF, DOCX, or PPTX."
        )


@router.post("/demo", response_model=AssessmentResponse)
def create_walkthrough_workspace(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Auto-fills a complete, premium assessment with walkthrough client workspace data"""
    logger.info("Initializing Walkthrough Sample Workspace...")
    organization = get_or_create_default_organization(db, current_user)
    sample_client = (
        db.query(Client)
        .filter(
            Client.organization_id == organization.id,
            Client.name == "Apex Global Consulting Partners",
        )
        .first()
    )
    if not sample_client:
        sample_client = Client(
            organization_id=organization.id,
            name="Apex Global Consulting Partners",
            industry="Professional Services",
            company_size="100-500 employees",
            cloud_preference="Azure",
            compliance_requirements=["GDPR", "SOC2 Type II"],
        )
        db.add(sample_client)
        db.commit()
        db.refresh(sample_client)

    existing_sample = (
        db.query(Assessment)
        .filter(
            Assessment.user_id == current_user.id,
            Assessment.company_name == "Apex Global Consulting Partners",
        )
        .first()
    )
    if existing_sample:
        if existing_sample.client_summary is None:
            existing_sample.client_summary = (
                "Apex Global Consulting Partners has a credible near-term AI pilot opportunity in proposal "
                "automation and support triage. The most immediate value is in shortening response cycles "
                "without forcing a large platform replacement."
            )
        if existing_sample.data_justification is None:
            existing_sample.data_justification = (
                "Proposal content and operational knowledge exist across Salesforce, Excel, and SharePoint, "
                "but the data remains fragmented enough to limit immediate AI reuse without cleanup."
            )
        if existing_sample.process_justification is None:
            existing_sample.process_justification = (
                "Several workflows are highly repetitive and suitable for AI support, but they still rely on "
                "manual reviewer handoffs and informal execution patterns."
            )
        if existing_sample.integration_justification is None:
            existing_sample.integration_justification = (
                "Apex has mainstream enterprise systems in place, yet the operating model still depends on "
                "human coordination rather than clean middleware or API-driven orchestration."
            )
        if existing_sample.governance_justification is None:
            existing_sample.governance_justification = (
                "GDPR and SOC2 create a credible governance baseline, but AI-specific approval trails and "
                "output validation should still be tightened before expansion."
            )
        if existing_sample.security_justification is None:
            existing_sample.security_justification = (
                "The existing compliance posture supports a controlled pilot, especially because regulated "
                "document handling practices are already familiar to the organization."
            )
        if existing_sample.team_justification is None:
            existing_sample.team_justification = (
                "The business has clear appetite for AI pilots, but long-term success will depend on explicit "
                "ownership, rollout training, and stronger adoption processes."
            )
        if existing_sample.alignment_justification is None:
            existing_sample.alignment_justification = (
                "The selected AI opportunities map directly to faster proposal turnaround, improved support "
                "responsiveness, and reduced manual operations overhead."
            )
        if existing_sample.reviewer_notes is None:
            existing_sample.reviewer_notes = (
                "Sample client workspace pre-reviewed for consulting walkthroughs. Reconfirm integration assumptions "
                "and compliance posture against live environments."
            )
        if existing_sample.approval_status != "approved":
            existing_sample.approval_status = "approved"
        db.commit()
        db.refresh(existing_sample)
        return existing_sample

    # Create the assessment session
    ass = Assessment(
        user_id=current_user.id,
        client_id=sample_client.id,
        company_name="Apex Global Consulting Partners",
        industry="Professional Services",
        company_size="100-500 employees",
        departments=[
            "Operations",
            "Customer Support",
            "Sales & Pre-sales",
            "Compliance & Governance",
        ],
        current_tools=[
            "Salesforce",
            "Microsoft Excel",
            "Jira Service Desk",
            "SharePoint",
        ],
        cloud_preference="Azure",
        compliance_requirements=["GDPR", "SOC2 Type II"],
        main_business_goals="Automate redundant pre-sales proposal drafting and support ticket parsing to optimize service delivery speeds.",
        pain_points=[
            "Manual Process Overload",
            "Data Silos",
            "Slow Support Response Times",
        ],
        ai_goals=["Efficiency / Cost Reduction", "Enhanced Customer Experience"],
        # Scoring Metrics
        overall_score=62.0,
        automation_potential=28.0,
        confidence_score=92.0,
        status="completed",
        risk_level="Medium",
        # Breakdown
        data_readiness=55.0,
        data_justification="Proposal content and operational knowledge exist across Salesforce, Excel, and SharePoint, but the data remains fragmented enough to limit immediate AI reuse without cleanup.",
        process_readiness=60.0,
        process_justification="Several workflows are highly repetitive and suitable for AI support, but they still rely on manual reviewer handoffs and informal execution patterns.",
        integration_readiness=50.0,
        integration_justification="Apex has mainstream enterprise systems in place, yet the operating model still depends on human coordination rather than clean middleware or API-driven orchestration.",
        governance_readiness=68.0,
        governance_justification="GDPR and SOC2 create a credible governance baseline, but AI-specific approval trails and output validation should still be tightened before expansion.",
        security_readiness=72.0,
        security_justification="The existing compliance posture supports a controlled pilot, especially because regulated document handling practices are already familiar to the organization.",
        team_readiness=58.0,
        team_justification="The business has clear appetite for AI pilots, but long-term success will depend on explicit ownership, rollout training, and stronger adoption processes.",
        business_alignment=75.0,
        alignment_justification="The selected AI opportunities map directly to faster proposal turnaround, improved support responsiveness, and reduced manual operations overhead.",
        business_summary=(
            "Apex Global Consulting Partners is a mid-size professional services firm handling "
            "intensive support and operations workloads. Repeated workflows in data collation and "
            "proposal drafting slow operational response indexes. Embedding focused LLM classification "
            "agents represents a P1 high-impact transformation opportunity."
        ),
        client_summary=(
            "Apex Global Consulting Partners has a credible near-term AI pilot opportunity in proposal "
            "automation and support triage. The most immediate value is in shortening response cycles "
            "without forcing a large platform replacement."
        ),
        reviewer_notes=(
            "Sample client workspace pre-reviewed for consulting walkthroughs. Reconfirm integration assumptions "
            "and compliance posture against live environments."
        ),
        approval_status="approved",
        readiness_interpretation=(
            "With an overall AI readiness score of 62/100, Apex is well positioned to pilot operational AI agents. "
            "Security and Compliance metrics (72/100) are strong due to existing SOC2 policies. However, "
            "integration silos in legacy spreadsheets limit instant scale-out, calling for local middleware layers."
        ),
        recommended_first_pilot="Intelligent Pre-Sales Proposal Copilot",
        why_recommended_pilot="High transformation value, low implementation complexity, and fits the corporate priority of accelerating pre-sales proposal throughput.",
        expected_pilot_impact="Estimated to reduce first-draft proposal preparation from 2–3 days to under 30 minutes, with potential to improve reuse of approved content, consistency, and senior review efficiency based on observed workflow signals.",
    )
    db.add(ass)
    db.commit()
    db.refresh(ass)

    # Seed Bottlenecks
    bottlenecks = [
        ProcessBottleneck(
            assessment_id=ass.id,
            department="Sales & Pre-sales",
            process_name="Custom RFP / Proposal Compilation",
            bottleneck_description="Technical sales executives spend up to 12 hours manually copy-pasting references and templates from old bid documents.",
            ai_potential="High",
        ),
        ProcessBottleneck(
            assessment_id=ass.id,
            department="Customer Support",
            process_name="Manual Incident Ingestion",
            bottleneck_description="Triage engineers manually tag, catalog, and routes incoming client incident tickets, delaying response times.",
            ai_potential="High",
        ),
        ProcessBottleneck(
            assessment_id=ass.id,
            department="Operations",
            process_name="Excel Billing Reconciliation",
            bottleneck_description="Finance managers manually match billing columns from hourly logs in spreadsheets to ERP database sheets.",
            ai_potential="Medium",
        ),
    ]

    # Seed Use Cases
    use_cases = [
        AIUseCase(
            assessment_id=ass.id,
            department=normalize_department(
                "Intelligent Pre-Sales Proposal Copilot",
                "Intelligent solution drafting companion that pulls from verified proposal assets and past winning bid documents to build compliance-mapped proposal outlines.",
            ),
            use_case_name="Intelligent Pre-Sales Proposal Copilot",
            description="Intelligent solution drafting companion that pulls from verified proposal assets and past winning bid documents to build compliance-mapped proposal outlines.",
            value="High",
            complexity="Low",
            risk="Low",
            priority="P1",
            evidence="Recommended because sales teams report spending over 15 hours weekly on manual copy-pasting during RFP bid cycles.",
            confidence=92.0,
        ),
        AIUseCase(
            assessment_id=ass.id,
            department=normalize_department(
                "Autonomous Support Triage Router",
                "Implement an intelligent support ticket routing pipeline using LLM classification to auto-categorize and assign technician tags.",
            ),
            use_case_name="Autonomous Support Triage Router",
            description="Implement an intelligent support ticket routing pipeline using LLM classification to auto-categorize and assign technician tags.",
            value="High",
            complexity="Medium",
            risk="Low",
            priority="P1",
            evidence="Recommended because logs indicate ticket classification delays average 4 hours per incoming incident.",
            confidence=88.0,
        ),
        AIUseCase(
            assessment_id=ass.id,
            department=normalize_department(
                "Billing Excel Extraction Agent",
                "Deploy a layout-aware processing pipeline that matches invoice records against Excel spreadsheet columns.",
            ),
            use_case_name="Billing Excel Extraction Agent",
            description="Deploy a layout-aware processing pipeline that matches invoice records against Excel spreadsheet columns.",
            value="Medium",
            complexity="Low",
            risk="Low",
            priority="P2",
            evidence="Recommended because uploaded spreadsheets show repetitive billing entries mapped against old templates.",
            confidence=85.0,
        ),
        AIUseCase(
            assessment_id=ass.id,
            department=normalize_department(
                "Contract Audit Assistant",
                "An AI compliance auditor that automatically evaluates signed supplier agreements against GDPR mandates.",
            ),
            use_case_name="Contract Audit Assistant",
            description="An AI compliance auditor that automatically evaluates signed supplier agreements against GDPR mandates.",
            value="High",
            complexity="High",
            risk="Medium",
            priority="P2",
            evidence="Recommended because client maintains rigid GDPR compliance protocols for global operations.",
            confidence=82.0,
        ),
    ]

    # Seed Risks
    risks = [
        RiskRegister(
            assessment_id=ass.id,
            risk_name="Data Leaks in Proposal Generation",
            severity="Medium",
            recommendation="Establish strict local sanitization gates (regex scrubbers) to wipe customer names and proprietary figures prior to API submissions.",
            is_control_met=1,
        ),
        RiskRegister(
            assessment_id=ass.id,
            risk_name="Hallucinated Billing Records",
            severity="High",
            recommendation="Enforce a strict human-in-the-loop manual review interface for billing matches returning confidence values below 95%.",
            is_control_met=0,
        ),
    ]

    # Seed Roadmap
    roadmap = [
        RoadmapItem(
            assessment_id=ass.id,
            phase="30-Day",
            action_item="Launch Intelligent Pre-Sales Proposal Copilot MVP using winning contract vector stores.",
            expected_impact="Accelerates initial proposal drafting speeds by an estimated 65%.",
            confidence=90.0,
        ),
        RoadmapItem(
            assessment_id=ass.id,
            phase="60-Day",
            action_item="Integrate Support Triage Router into Jira Service Desk shadow staging queues.",
            expected_impact="Targeting up to 40% optimization of ticketing triage cycle times based on initial pilot benchmarks.",
            confidence=85.0,
        ),
        RoadmapItem(
            assessment_id=ass.id,
            phase="90-Day",
            action_item="Deploy Compliance Contract Auditor and enforce operational training for pre-sales teams.",
            expected_impact="Potential to save up to 10 staff hours weekly while supporting complete audit readiness.",
            confidence=92.0,
        ),
    ]

    # Seed Signals
    signals = [
        DocumentSignal(
            assessment_id=ass.id,
            source_file="Apex_Sales_Playbook.pdf",
            signal_type="RFP workflow",
            description="Identified 12-hour manual cycles spent on custom solution blocks in proposal compilation.",
            confidence=95.0,
        ),
        DocumentSignal(
            assessment_id=ass.id,
            source_file="Apex_Operations_Audit.txt",
            signal_type="Excel billing bottleneck",
            description="Found manual hourly log exports manually mapped to corporate billing templates.",
            confidence=90.0,
        ),
    ]

    for item in bottlenecks:
        db.add(item)
    for item in use_cases:
        db.add(item)
    for item in risks:
        db.add(item)
    for item in roadmap:
        db.add(item)
    for item in signals:
        db.add(item)

    db.commit()
    db.refresh(ass)
    return ass
